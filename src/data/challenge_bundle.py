"""Helpers for working with the official India Runs public challenge bundle."""

from __future__ import annotations

import gzip
import json
import re
from collections.abc import Iterable, Iterator
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document

from src.utils.paths import DATA_DIR, PROCESSED_DATA_DIR, RAW_DATA_DIR

TRACK1_JOB_ID = "REDROB_TRACK1_MAIN_JD"
CANONICAL_JOB_DATASET_NAME = "challenge_jobs.csv"
CANONICAL_CANDIDATE_DATASET_NAME = "challenge_candidates.parquet"

SECTION_HEADINGS = (
    "Let's be honest about this role",
    "What you'd actually be doing",
    'What we mean by "5-9 years"',
    "The skills inventory (please read carefully)",
    "Things you absolutely need",
    "Things we'd like you to have but won't reject you for",
    "Things we explicitly do NOT want",
    "On location, comp, and logistics",
    "The vibe check",
    "How to read between the lines",
    "Final note for the participants of the Redrob hackathon",
)

EXPERIENCE_RANGE_PATTERN = re.compile(
    r"(?P<minimum>\d+(?:\.\d+)?)\s*[–-]\s*(?P<maximum>\d+(?:\.\d+)?)\s+years",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class ChallengeBundlePaths:
    """Resolved filesystem paths for the public challenge bundle."""

    root: Path
    candidates: Path
    candidate_schema: Path
    job_description: Path
    bundle_readme: Path
    redrob_signals: Path
    submission_spec: Path
    sample_submission: Path


def discover_challenge_bundle(search_dirs: Iterable[Path] | None = None) -> ChallengeBundlePaths:
    """Locate the extracted challenge bundle on disk."""
    roots = tuple(search_dirs or (RAW_DATA_DIR, DATA_DIR))
    candidate_names = ("candidates.jsonl", "candidates.jsonl.gz")

    for root in roots:
        if not root.exists():
            continue
        for filename in candidate_names:
            for candidate_path in sorted(root.rglob(filename)):
                bundle_root = candidate_path.parent
                job_description = bundle_root / "job_description.docx"
                candidate_schema = bundle_root / "candidate_schema.json"
                sample_submission = bundle_root / "sample_submission.csv"
                submission_spec = bundle_root / "submission_spec.docx"
                bundle_readme = bundle_root / "README.docx"
                redrob_signals = bundle_root / "redrob_signals_doc.docx"
                if all(
                    path.exists()
                    for path in (
                        job_description,
                        candidate_schema,
                        sample_submission,
                        submission_spec,
                        bundle_readme,
                        redrob_signals,
                    )
                ):
                    return ChallengeBundlePaths(
                        root=bundle_root,
                        candidates=candidate_path,
                        candidate_schema=candidate_schema,
                        job_description=job_description,
                        bundle_readme=bundle_readme,
                        redrob_signals=redrob_signals,
                        submission_spec=submission_spec,
                        sample_submission=sample_submission,
                    )

    raise FileNotFoundError(
        "Could not locate the extracted India Runs challenge bundle under data/raw or data/."
    )


def read_docx_text(path: str | Path) -> str:
    """Extract paragraph and table text from a DOCX file."""
    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cell for cell in cells if cell))

    return "\n".join(blocks)


def _sectionise_lines(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current_heading = "Preamble"
    buffer: list[str] = []

    for line in lines:
        if line in SECTION_HEADINGS:
            sections[current_heading] = buffer
            current_heading = line
            buffer = []
            continue
        buffer.append(line)

    sections[current_heading] = buffer
    return {heading: values for heading, values in sections.items() if values}


def parse_job_description(text: str) -> dict[str, object]:
    """Parse the released job-description document into a tabular row."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    sections = _sectionise_lines(lines)

    metadata_map = {
        "Job Description:": "job_title",
        "Company:": "company",
        "Location:": "location",
        "Employment Type:": "employment_type",
        "Experience Required:": "experience_required_text",
    }
    metadata: dict[str, str] = {}
    for line in lines[:8]:
        for prefix, key in metadata_map.items():
            if line.startswith(prefix):
                metadata[key] = line.removeprefix(prefix).strip()

    experience_text = metadata.get("experience_required_text", "")
    minimum_experience: float | None = None
    maximum_experience: float | None = None
    match = EXPERIENCE_RANGE_PATTERN.search(experience_text)
    if match:
        minimum_experience = float(match.group("minimum"))
        maximum_experience = float(match.group("maximum"))

    def join_section(name: str) -> str:
        return "\n".join(sections.get(name, [])).strip()

    return {
        "job_id": TRACK1_JOB_ID,
        "job_title": metadata.get("job_title", "Senior AI Engineer - Founding Team"),
        "company": metadata.get("company", ""),
        "location": metadata.get("location", ""),
        "employment_type": metadata.get("employment_type", ""),
        "experience_required_text": experience_text,
        "min_experience": minimum_experience,
        "max_experience": maximum_experience,
        "required_skills": join_section("Things you absolutely need"),
        "preferred_skills": join_section("Things we'd like you to have but won't reject you for"),
        "excluded_profiles": join_section("Things we explicitly do NOT want"),
        "responsibilities": join_section("What you'd actually be doing"),
        "ideal_candidate": join_section("How to read between the lines"),
        "job_description": text,
    }


def save_canonical_job_dataset(bundle: ChallengeBundlePaths) -> Path:
    """Create the flattened one-row jobs dataset used by the repo."""
    job_row = parse_job_description(read_docx_text(bundle.job_description))
    PROCESSED_DATA_DIR.mkdir(parents=True, exist_ok=True)
    output_path = PROCESSED_DATA_DIR / CANONICAL_JOB_DATASET_NAME
    pd.DataFrame([job_row]).to_csv(output_path, index=False)
    return output_path


def stream_candidate_records(
    path: str | Path,
    *,
    limit: int | None = None,
) -> Iterator[dict[str, object]]:
    """Yield candidate records from the released JSONL or JSONL.GZ file."""
    path = Path(path)
    opener = gzip.open if path.suffixes[-2:] == [".jsonl", ".gz"] else open
    open_kwargs = {"mode": "rt", "encoding": "utf-8"}

    with opener(path, **open_kwargs) as handle:
        for index, line in enumerate(handle):
            if limit is not None and index >= limit:
                break
            payload = line.strip()
            if payload:
                yield json.loads(payload)


def _clean_join(parts: Iterable[str], *, separator: str = " | ") -> str:
    return separator.join(part.strip() for part in parts if part and part.strip())


def flatten_candidate_record(record: dict[str, object]) -> dict[str, object]:
    """Flatten one nested candidate profile into a table-friendly row."""
    profile = record.get("profile") or {}
    career_history = record.get("career_history") or []
    education = record.get("education") or []
    skills = record.get("skills") or []
    certifications = record.get("certifications") or []
    languages = record.get("languages") or []
    signals = record.get("redrob_signals") or {}

    skill_names = [str(item.get("name", "")).strip() for item in skills if item.get("name")]
    skill_details = [
        _clean_join(
            (
                str(item.get("name", "")).strip(),
                str(item.get("proficiency", "")).strip(),
                f"endorsements={item.get('endorsements', 0)}",
                f"months={item.get('duration_months', 0)}",
            ),
            separator="; ",
        )
        for item in skills
        if item.get("name")
    ]
    career_history_text = _clean_join(
        [
            _clean_join(
                (
                    str(entry.get("title", "")).strip(),
                    f"at {str(entry.get('company', '')).strip()}".strip(),
                    str(entry.get("description", "")).strip(),
                ),
                separator=". ",
            )
            for entry in career_history
            if any(entry.get(field) for field in ("title", "company", "description"))
        ],
        separator=" || ",
    )
    education_text = _clean_join(
        [
            _clean_join(
                (
                    str(item.get("degree", "")).strip(),
                    str(item.get("field_of_study", "")).strip(),
                    str(item.get("institution", "")).strip(),
                ),
                separator=", ",
            )
            for item in education
            if any(item.get(field) for field in ("degree", "field_of_study", "institution"))
        ]
    )
    certifications_text = _clean_join(
        [
            _clean_join(
                (
                    str(item.get("name", "")).strip(),
                    str(item.get("issuer", "")).strip(),
                    str(item.get("year", "")).strip(),
                ),
                separator=", ",
            )
            for item in certifications
            if any(item.get(field) for field in ("name", "issuer", "year"))
        ]
    )
    languages_text = _clean_join(
        [
            _clean_join(
                (
                    str(item.get("language", "")).strip(),
                    str(item.get("proficiency", "")).strip(),
                ),
                separator=", ",
            )
            for item in languages
            if any(item.get(field) for field in ("language", "proficiency"))
        ]
    )

    assessment_scores = signals.get("skill_assessment_scores") or {}
    assessment_values = [
        float(value)
        for value in assessment_scores.values()
        if isinstance(value, int | float)
    ]
    expected_salary = signals.get("expected_salary_range_inr_lpa") or {}

    candidate_text = _clean_join(
        (
            str(profile.get("headline", "")).strip(),
            str(profile.get("summary", "")).strip(),
            str(profile.get("current_title", "")).strip(),
            str(profile.get("current_company", "")).strip(),
            str(profile.get("current_industry", "")).strip(),
            ", ".join(skill_names),
            career_history_text,
            education_text,
            certifications_text,
            languages_text,
        )
    )

    return {
        "candidate_id": record.get("candidate_id"),
        "current_role": profile.get("current_title"),
        "headline": profile.get("headline"),
        "summary": profile.get("summary"),
        "location": profile.get("location"),
        "country": profile.get("country"),
        "current_company": profile.get("current_company"),
        "current_company_size": profile.get("current_company_size"),
        "current_industry": profile.get("current_industry"),
        "total_experience": profile.get("years_of_experience"),
        "skills": ", ".join(skill_names),
        "skills_detailed": " || ".join(detail for detail in skill_details if detail),
        "education": education_text,
        "career_history_text": career_history_text,
        "certifications": certifications_text,
        "languages": languages_text,
        "profile_text": candidate_text,
        "profile_completeness_score": signals.get("profile_completeness_score"),
        "signup_date": signals.get("signup_date"),
        "last_active": signals.get("last_active_date"),
        "open_to_work_flag": signals.get("open_to_work_flag"),
        "profile_views_received_30d": signals.get("profile_views_received_30d"),
        "applications_submitted_30d": signals.get("applications_submitted_30d"),
        "recruiter_response_rate": signals.get("recruiter_response_rate"),
        "avg_response_time_hours": signals.get("avg_response_time_hours"),
        "connection_count": signals.get("connection_count"),
        "endorsements_received": signals.get("endorsements_received"),
        "notice_period_days": signals.get("notice_period_days"),
        "expected_salary_min_inr_lpa": expected_salary.get("min"),
        "expected_salary_max_inr_lpa": expected_salary.get("max"),
        "preferred_work_mode": signals.get("preferred_work_mode"),
        "willing_to_relocate": signals.get("willing_to_relocate"),
        "github_activity_score": signals.get("github_activity_score"),
        "search_appearance_30d": signals.get("search_appearance_30d"),
        "saved_by_recruiters_30d": signals.get("saved_by_recruiters_30d"),
        "interview_completion_rate": signals.get("interview_completion_rate"),
        "offer_acceptance_rate": signals.get("offer_acceptance_rate"),
        "verified_email": signals.get("verified_email"),
        "verified_phone": signals.get("verified_phone"),
        "linkedin_connected": signals.get("linkedin_connected"),
        "skill_count": len(skill_names),
        "skill_assessment_count": len(assessment_scores),
        "skill_assessment_avg": (
            sum(assessment_values) / len(assessment_values) if assessment_values else None
        ),
        "career_history_count": len(career_history),
        "education_count": len(education),
        "certification_count": len(certifications),
        "language_count": len(languages),
    }


def save_canonical_candidate_dataset(bundle: ChallengeBundlePaths) -> Path:
    """Create the flattened candidate table used by the repo."""
    rows = [flatten_candidate_record(record) for record in stream_candidate_records(bundle.candidates)]
    PROCESSED_DATA_DIR.mkdir(parents=True, exist_ok=True)
    output_path = PROCESSED_DATA_DIR / CANONICAL_CANDIDATE_DATASET_NAME
    pd.DataFrame.from_records(rows).to_parquet(output_path, index=False)
    return output_path


def load_candidate_id_set(bundle: ChallengeBundlePaths) -> set[str]:
    """Load the full released candidate ID set for submission validation."""
    return {
        str(record.get("candidate_id"))
        for record in stream_candidate_records(bundle.candidates)
        if record.get("candidate_id")
    }
